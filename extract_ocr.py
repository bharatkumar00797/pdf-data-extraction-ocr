import os
from pdf2image import convert_from_path
import pytesseract
from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Font
from datetime import datetime

# ====================== SETTINGS ======================
PDF_FOLDER = r"C:\Users\padma\OneDrive\Desktop\Anjuman\her office work august 28\Annual reports-20260804T103947Z-1-001\Annual reports"

# >>> CHANGE THESE TWO PATHS according to your installation <<<
TESSERACT_PATH = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
POPPLER_PATH = r"C:\poppler\Library\bin"

OUTPUT_WORD = "OCR_Extracted_Reports.docx"
OUTPUT_EXCEL = "OCR_Extracted_Reports.xlsx"

# Optional: process only one file first for testing (set to None to process all)
TEST_ONLY_ONE_FILE = None   # example: "15th  English annual report .pdf"
# ======================================================

pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH

def main():
    print("=" * 65)
    print("OCR Extraction Started (this will take time)")
    print("=" * 65)

    if not os.path.exists(PDF_FOLDER):
        print("ERROR: PDF folder not found")
        return

    pdf_files = [f for f in os.listdir(PDF_FOLDER) if f.lower().endswith(".pdf")]
    pdf_files.sort()

    if TEST_ONLY_ONE_FILE:
        pdf_files = [f for f in pdf_files if f == TEST_ONLY_ONE_FILE]
        print(f"Testing mode: only processing → {TEST_ONLY_ONE_FILE}\n")
    else:
        print(f"Found {len(pdf_files)} PDF files. This will take a long time...\n")

    doc = Document()
    doc.add_heading("OCR Extracted Text from Annual Reports", level=0)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%d %B %Y, %I:%M %p')}")
    doc.add_paragraph("Note: Text extracted using OCR from scanned images. Accuracy is not 100%.")
    doc.add_paragraph("")

    wb = Workbook()
    default = wb.active
    wb.remove(default)

    for pdf_file in pdf_files:
        pdf_path = os.path.join(PDF_FOLDER, pdf_file)
        print(f"\nProcessing: {pdf_file}")

        doc.add_heading(pdf_file, level=1)

        sheet_name = pdf_file.replace(".pdf", "").replace(".PDF", "")[:31]
        ws = wb.create_sheet(title=sheet_name)
        ws["A1"] = "Page"
        ws["B1"] = "OCR Text"
        ws["A1"].font = Font(bold=True)
        ws["B1"].font = Font(bold=True)
        ws.column_dimensions["A"].width = 8
        ws.column_dimensions["B"].width = 120

        try:
            print("   Converting pages to images...")
            images = convert_from_path(
                pdf_path,
                dpi=200,
                poppler_path=POPPLER_PATH
            )

            for page_num, image in enumerate(images, start=1):
                print(f"   OCR on page {page_num}/{len(images)}...", end="\r")

                text = pytesseract.image_to_string(image, lang="eng")

                doc.add_heading(f"Page {page_num}", level=2)
                if text.strip():
                    para = doc.add_paragraph(text.strip())
                    para.paragraph_format.space_after = Pt(6)
                else:
                    doc.add_paragraph("[No text detected]")

                ws.cell(row=page_num + 1, column=1, value=page_num)
                ws.cell(row=page_num + 1, column=2, value=text)

            print(f"\n   → Finished {len(images)} pages")

        except Exception as e:
            print(f"\n   ERROR: {e}")
            doc.add_paragraph(f"Error processing this file: {e}")

        doc.add_page_break()

    doc.save(OUTPUT_WORD)
    wb.save(OUTPUT_EXCEL)

    print("\n" + "=" * 65)
    print("OCR EXTRACTION COMPLETED")
    print("=" * 65)
    print(f"Word file : {OUTPUT_WORD}")
    print(f"Excel file: {OUTPUT_EXCEL}")


if __name__ == "__main__":
    main()
