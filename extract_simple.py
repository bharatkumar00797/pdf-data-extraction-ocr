import os
import pdfplumber
from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Font
from datetime import datetime

# ====================== FOLDER PATH ======================
PDF_FOLDER = r"C:\Users\padma\OneDrive\Desktop\Anjuman\her office work august 28\Annual reports-20260804T103947Z-1-001\Annual reports"

OUTPUT_WORD = "All_Reports_Extracted.docx"
OUTPUT_EXCEL = "All_Reports_Extracted.xlsx"
# =========================================================

def main():
    print("=" * 60)
    print("Simple Text Extraction Started")
    print("=" * 60)

    if not os.path.exists(PDF_FOLDER):
        print("ERROR: Folder not found!")
        print(PDF_FOLDER)
        return

    pdf_files = [f for f in os.listdir(PDF_FOLDER) if f.lower().endswith(".pdf")]
    pdf_files.sort()

    print(f"Found {len(pdf_files)} PDF files\n")

    # ---------- Create Word Document ----------
    doc = Document()
    doc.add_heading("All Annual Reports - Extracted Text", level=0)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%d %B %Y, %I:%M %p')}")
    doc.add_paragraph("")

    # ---------- Create Excel Workbook ----------
    wb = Workbook()
    default_sheet = wb.active
    wb.remove(default_sheet)

    for pdf_file in pdf_files:
        pdf_path = os.path.join(PDF_FOLDER, pdf_file)
        print(f"Processing: {pdf_file}")

        doc.add_heading(pdf_file, level=1)

        sheet_name = pdf_file.replace(".pdf", "").replace(".PDF", "")[:31]
        ws = wb.create_sheet(title=sheet_name)
        ws["A1"] = "Page"
        ws["B1"] = "Extracted Text"
        ws["A1"].font = Font(bold=True)
        ws["B1"].font = Font(bold=True)
        ws.column_dimensions["A"].width = 10
        ws.column_dimensions["B"].width = 120

        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    text = page.extract_text()
                    if not text:
                        text = ""

                    doc.add_heading(f"Page {page_num}", level=2)
                    if text.strip():
                        para = doc.add_paragraph(text.strip())
                        para.paragraph_format.space_after = Pt(8)
                    else:
                        doc.add_paragraph("[No text found on this page]")

                    ws.cell(row=page_num + 1, column=1, value=page_num)
                    ws.cell(row=page_num + 1, column=2, value=text)

            print(f"   → Done ({len(pdf.pages)} pages)")

        except Exception as e:
            print(f"   → Error: {e}")
            doc.add_paragraph(f"Error reading this file: {e}")

        doc.add_page_break()

    doc.save(OUTPUT_WORD)
    wb.save(OUTPUT_EXCEL)

    print("\n" + "=" * 60)
    print("FINISHED SUCCESSFULLY")
    print("=" * 60)
    print(f"Word file : {OUTPUT_WORD}")
    print(f"Excel file: {OUTPUT_EXCEL}")


if __name__ == "__main__":
    main()
