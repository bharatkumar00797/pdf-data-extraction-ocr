import os
import re
from datetime import datetime
from pdf2image import convert_from_path
import pytesseract
import pdfplumber
from docx import Document
from docx.shared import Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Font

# ====================== PATHS ======================
SCANNED_PDF_FOLDER = r"C:\Users\padma\OneDrive\Desktop\Anjuman\her office work august 28\Annual reports-20260804T103947Z-1-001\Annual reports"
EXCEL_PDF_PATH = r"C:\Users\padma\Downloads\divyajyoti-historical-dataset-v9.pdf"

# >>> EDIT THESE TWO PATHS <<<
TESSERACT_PATH = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
POPPLER_PATH   = r"C:\poppler\Library\bin"

OUTPUT_FOLDER = r"C:\Users\padma\Documents\OCR_Comparison_Results"
# ===================================================

pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH

os.makedirs(OUTPUT_FOLDER, exist_ok=True)
OCR_TEXT_FOLDER = os.path.join(OUTPUT_FOLDER, "OCR_Texts")
os.makedirs(OCR_TEXT_FOLDER, exist_ok=True)


def ocr_one_pdf(pdf_path, output_txt_path):
    """OCR a single scanned PDF and save text"""
    print(f"   Converting to images...")
    images = convert_from_path(pdf_path, dpi=200, poppler_path=POPPLER_PATH)
    
    all_text = []
    for i, img in enumerate(images, 1):
        print(f"   OCR page {i}/{len(images)}", end="\r")
        text = pytesseract.image_to_string(img, lang="eng")
        all_text.append(f"\n\n----- PAGE {i} -----\n\n{text}")
    
    full_text = "".join(all_text)
    with open(output_txt_path, "w", encoding="utf-8") as f:
        f.write(full_text)
    print(f"\n   Saved: {os.path.basename(output_txt_path)}")
    return full_text


def extract_text_from_clean_pdf(pdf_path):
    """Extract text from the Excel-converted (clean) PDF"""
    text_pages = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text_pages.append(t)
    return "\n\n".join(text_pages)


def extract_numbers(text):
    """Extract all numbers (including decimals and numbers with commas)"""
    pattern = r'\b\d{1,3}(?:,\d{3})*(?:\.\d+)?\b|\b\d+\.\d+\b|\b\d+\b'
    numbers = re.findall(pattern, text)
    cleaned = []
    for n in numbers:
        n_clean = n.replace(",", "")
        try:
            float(n_clean)
            cleaned.append(n_clean)
        except:
            pass
    return cleaned


def main():
    print("=" * 70)
    print("COMPARISON: Excel PDF  vs  15 Scanned Annual Reports")
    print("=" * 70)
    print(f"Start time: {datetime.now().strftime('%d-%b-%Y %I:%M %p')}\n")

    # ---------- Stage 1: OCR all 15 scanned PDFs ----------
    print("STAGE 1: OCR of 15 scanned PDFs")
    print("-" * 50)

    scanned_files = [f for f in os.listdir(SCANNED_PDF_FOLDER) if f.lower().endswith(".pdf")]
    scanned_files.sort()
    print(f"Found {len(scanned_files)} scanned PDFs\n")

    all_ocr_text = ""

    for idx, pdf_name in enumerate(scanned_files, 1):
        pdf_path = os.path.join(SCANNED_PDF_FOLDER, pdf_name)
        txt_name = pdf_name.rsplit(".", 1)[0] + ".txt"
        txt_path = os.path.join(OCR_TEXT_FOLDER, txt_name)

        print(f"[{idx}/{len(scanned_files)}] {pdf_name}")

        if os.path.exists(txt_path):
            print("   Already OCR'ed earlier → loading existing text")
            with open(txt_path, "r", encoding="utf-8") as f:
                text = f.read()
        else:
            try:
                text = ocr_one_pdf(pdf_path, txt_path)
            except Exception as e:
                print(f"   ERROR: {e}")
                text = ""
        
        all_ocr_text += f"\n\n========== {pdf_name} ==========\n\n{text}"

    # Save combined OCR text
    combined_path = os.path.join(OUTPUT_FOLDER, "ALL_15_REPORTS_OCR_COMBINED.txt")
    with open(combined_path, "w", encoding="utf-8") as f:
        f.write(all_ocr_text)
    print(f"\nCombined OCR text saved: {combined_path}")

    # ---------- Stage 2: Extract Excel PDF ----------
    print("\nSTAGE 2: Extracting text from Excel-converted PDF")
    print("-" * 50)

    excel_text = extract_text_from_clean_pdf(EXCEL_PDF_PATH)
    excel_txt_path = os.path.join(OUTPUT_FOLDER, "EXCEL_PDF_TEXT.txt")
    with open(excel_txt_path, "w", encoding="utf-8") as f:
        f.write(excel_text)
    print(f"Excel PDF text saved: {excel_txt_path}")

    # ---------- Stage 3: Comparison ----------
    print("\nSTAGE 3: Comparing data...")
    print("-" * 50)

    excel_numbers = extract_numbers(excel_text)
    unique_excel_numbers = list(dict.fromkeys(excel_numbers))

    print(f"Found {len(unique_excel_numbers)} unique numbers in Excel PDF")

    found = []
    not_found = []

    for num in unique_excel_numbers:
        if num in all_ocr_text:
            found.append(num)
        else:
            not_found.append(num)

    accuracy = (len(found) / len(unique_excel_numbers) * 100) if unique_excel_numbers else 0

    # ---------- Create Report ----------
    print("\nCreating final report...")

    doc = Document()
    doc.add_heading("Comparison Report: Excel PDF vs 15 Scanned Reports", 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%d %B %Y, %I:%M %p')}")
    doc.add_paragraph("")

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(f"Total unique numbers found in Excel PDF : {len(unique_excel_numbers)}")
    doc.add_paragraph(f"Numbers found in the 15 scanned reports : {len(found)}")
    doc.add_paragraph(f"Numbers NOT found                       : {len(not_found)}")
    doc.add_paragraph(f"Approximate Match Rate                  : {accuracy:.1f}%")
    doc.add_paragraph("")
    doc.add_paragraph("Note: This is an approximate check based on OCR. OCR can miss or misread numbers.")

    doc.add_heading("Numbers FOUND in original reports", level=1)
    if found:
        doc.add_paragraph(", ".join(found[:200]) + (" ..." if len(found) > 200 else ""))
    else:
        doc.add_paragraph("None")

    doc.add_heading("Numbers NOT FOUND in original reports", level=1)
    if not_found:
        para = doc.add_paragraph(", ".join(not_found[:200]) + (" ..." if len(not_found) > 200 else ""))
        for run in para.runs:
            run.font.color.rgb = RGBColor(200, 0, 0)
    else:
        doc.add_paragraph("None – All numbers were found")

    word_path = os.path.join(OUTPUT_FOLDER, "Comparison_Report.docx")
    doc.save(word_path)

    wb = Workbook()
    ws = wb.active
    ws.title = "Summary"
    ws["A1"] = "Comparison Summary"
    ws["A1"].font = Font(bold=True, size=14)
    ws["A3"] = "Total unique numbers in Excel PDF"
    ws["B3"] = len(unique_excel_numbers)
    ws["A4"] = "Numbers found in scanned reports"
    ws["B4"] = len(found)
    ws["A5"] = "Numbers NOT found"
    ws["B5"] = len(not_found)
    ws["A6"] = "Approximate Match Rate (%)"
    ws["B6"] = round(accuracy, 1)

    ws2 = wb.create_sheet("Found Numbers")
    ws2["A1"] = "Number"
    ws2["A1"].font = Font(bold=True)
    for i, n in enumerate(found, 2):
        ws2[f"A{i}"] = n

    ws3 = wb.create_sheet("Not Found Numbers")
    ws3["A1"] = "Number"
    ws3["A1"].font = Font(bold=True)
    for i, n in enumerate(not_found, 2):
        ws3[f"A{i}"] = n

    excel_report_path = os.path.join(OUTPUT_FOLDER, "Comparison_Report.xlsx")
    wb.save(excel_report_path)

    print("\n" + "=" * 70)
    print("COMPLETED")
    print("=" * 70)
    print(f"Results folder : {OUTPUT_FOLDER}")
    print(f"Word Report    : Comparison_Report.docx")
    print(f"Excel Report   : Comparison_Report.xlsx")
    print(f"Match Rate     : {accuracy:.1f}%")
    print("=" * 70)


if __name__ == "__main__":
    main()
